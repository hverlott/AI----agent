import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_markdown_table(lines):
    """
    Simple parser for Markdown tables.
    Returns a list of rows, where each row is a list of cell content.
    """
    table_data = []
    for line in lines:
        if not line.strip().startswith('|'):
            continue
        # Remove leading/trailing pipes and split
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # Skip separator lines (e.g., ---|---|---)
        if all(re.match(r'^[\-\:]+$', c) for c in cells):
            continue
        table_data.append(cells)
    return table_data

def add_formatted_text(paragraph, text):
    """
    Parses **bold** text and adds runs to the paragraph.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def md_to_docx(md_file_path, docx_file_path):
    document = Document()
    
    # Set default style if needed, or just rely on 'Normal'
    style = document.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei' # Good for Chinese
    font.size = Pt(11)

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        
        # Horizontal Rule
        elif line.startswith('---') or line.startswith('***'):
            document.add_paragraph('_' * 40) # Simple visual separator

        # List Items
        elif line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        
        # Numbered List (Simple detection)
        elif re.match(r'^\d+\.\s', line):
            p = document.add_paragraph(style='List Number')
            # Remove the number prefix for the style to handle it, or keep it if style doesn't auto-number
            # Python-docx List Number style usually handles indentation but might not auto-increment correctly without complex xml.
            # So we keep the text as is but use the style for indentation.
            text_content = re.sub(r'^\d+\.\s', '', line)
            add_formatted_text(p, text_content)

        # Table detection
        elif line.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            table_data = parse_markdown_table(table_lines)
            if table_data:
                # Create table
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                table = document.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for r_idx, row_data in enumerate(table_data):
                    row = table.rows[r_idx]
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < len(row.cells):
                            # Handle bold in table cells too
                            p = row.cells[c_idx].paragraphs[0]
                            add_formatted_text(p, cell_text)
            
            continue # Already incremented i inside the inner loop

        # Normal Paragraph
        else:
            p = document.add_paragraph()
            add_formatted_text(p, line)

        i += 1

    document.save(docx_file_path)
    print(f"Successfully converted {md_file_path} to {docx_file_path}")

if __name__ == "__main__":
    target_md = r"d:\SaaS-AIs\releases\v2.5.1\docs\help_center\v1.0\zh_CN\客服AI系统说明文档.md"
    target_docx = r"d:\SaaS-AIs\releases\v2.5.1\docs\help_center\v1.0\zh_CN\客服AI系统说明文档.docx"
    md_to_docx(target_md, target_docx)
