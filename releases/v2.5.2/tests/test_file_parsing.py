import unittest
from unittest.mock import MagicMock, patch
import io
import sys

# Copied from admin_multi.py to verify logic without importing the whole streamlit app
# In a real refactor, this should be moved to a shared utility module.
def extract_content_from_upload(upload, filename):
    name_lower = (filename or "").lower()
    content = ""
    parse_note = ""
    if name_lower.endswith((".txt", ".md")):
        try:
            content = upload.getvalue().decode("utf-8", errors="ignore")
        except Exception:
            content = upload.getvalue().decode("latin-1", errors="ignore")
    elif name_lower.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(upload)
            pages = []
            for i in range(len(reader.pages)):
                text = reader.pages[i].extract_text()
                if text:
                    pages.append(f"# PDF Page {i+1}\n{text}")
            content = "\n\n".join(pages).strip()
            parse_note = "parsed:pdf"
        except Exception as e:
            parse_note = f"unparsed:pdf:{e}"
            content = ""
    elif name_lower.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(upload)
            content = "\n".join([p.text for p in doc.paragraphs]).strip()
            parse_note = "parsed:docx"
        except Exception as e:
            parse_note = f"unparsed:docx:{e}"
            content = ""
    elif name_lower.endswith(".xlsx"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(upload, read_only=True, data_only=True)
            texts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    items = [str(cell) for cell in row if cell is not None]
                    if items:
                        texts.append(" | ".join(items))
            content = "\n".join(texts).strip()
            parse_note = "parsed:xlsx"
        except Exception as e:
            parse_note = f"unparsed:xlsx:{e}"
            content = ""
    else:
        try:
            content = upload.getvalue().decode("utf-8", errors="ignore")
        except Exception:
            content = ""
        parse_note = "unknown_format"
    return content, parse_note

class TestFileParsing(unittest.TestCase):
    
    @patch('pypdf.PdfReader')
    def test_pdf_extract(self, mock_pdf_reader):
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "PDF Content"
        mock_reader_instance = MagicMock()
        mock_reader_instance.pages = [mock_page]
        mock_pdf_reader.return_value = mock_reader_instance
        
        mock_upload = io.BytesIO(b"fake pdf")
        content, note = extract_content_from_upload(mock_upload, "test.pdf")
        
        self.assertEqual(note, "parsed:pdf")
        self.assertIn("PDF Content", content)

    @patch('docx.Document')
    def test_docx_extract(self, mock_docx_doc):
        mock_para1 = MagicMock()
        mock_para1.text = "Para 1"
        mock_para2 = MagicMock()
        mock_para2.text = "Para 2"
        
        mock_doc_instance = MagicMock()
        mock_doc_instance.paragraphs = [mock_para1, mock_para2]
        mock_docx_doc.return_value = mock_doc_instance
        
        mock_upload = io.BytesIO(b"fake docx")
        content, note = extract_content_from_upload(mock_upload, "test.docx")
        
        self.assertEqual(note, "parsed:docx")
        self.assertIn("Para 1", content)
        self.assertIn("Para 2", content)

    @patch('openpyxl.load_workbook')
    def test_xlsx_extract(self, mock_load_wb):
        mock_ws = MagicMock()
        # Mock rows: [("A1", "B1"), ("A2", None)]
        mock_ws.iter_rows.return_value = [
            ("A1", "B1"),
            ("A2", None)
        ]
        
        mock_wb_instance = MagicMock()
        mock_wb_instance.worksheets = [mock_ws]
        mock_load_wb.return_value = mock_wb_instance
        
        mock_upload = io.BytesIO(b"fake xlsx")
        content, note = extract_content_from_upload(mock_upload, "test.xlsx")
        
        self.assertEqual(note, "parsed:xlsx")
        self.assertIn("A1 | B1", content)
        self.assertIn("A2", content)

    def test_txt_extract(self):
        mock_upload = io.BytesIO(b"Hello World")
        content, note = extract_content_from_upload(mock_upload, "test.txt")
        # Note: txt/md returns empty parse_note in current implementation unless error, 
        # but logic says parse_note = "" initially.
        self.assertEqual(content, "Hello World")

if __name__ == '__main__':
    unittest.main()
